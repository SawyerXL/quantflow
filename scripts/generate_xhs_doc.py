#!/usr/bin/env python3
"""Generate QuantFlow 小红书推广文案 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x24, 0x42)

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")

def tag(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x52, 0x52, 0x5B)

def divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—" * 30)
    run.font.color.rgb = RGBColor(0xA1, 0xA1, 0xAA)

# ═══════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("📱 小红书发布文案")
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("QuantFlow 推广素材 · 可直接复制发布").font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("发布地址：https://quantflow.pages.dev/demo").font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. 封面图说明
# ═══════════════════════════════════════════════════════════
h1("📸 封面图准备")

body("用手机打开 https://quantflow.pages.dev/demo")
body("→ 切到「BTC · Triple Moving Average」Tab")
body("→ 截图指标卡片区域（6 个数字卡 + 曲线图）")
body("→ 用美图秀秀或醒图加文字标注")
divider()

# ═══════════════════════════════════════════════════════════
# 2. 标题
# ═══════════════════════════════════════════════════════════
h1("标题（二选一）")

h2("方案 A（数据冲击型）：")
body("🚀 BTC 用这个策略，5 年收益 +1662%")

h2("方案 B（悬念型）：")
body("😱 回测了 12 种策略，BTC + Triple MA 炸了")
divider()

# ═══════════════════════════════════════════════════════════
# 3. 正文
# ═══════════════════════════════════════════════════════════
h1("正文（直接复制）")

body("🆓 先说重点：完全免费，不用写代码，浏览器打开就能用")
body("")
body("📊 我用自己做的回测工具 QuantFlow，把 BTC 过去 5 年的数据跑了 12 种策略。")
body("")
body("🏆 前三名：")
body("① Triple MA（三均线）：+1662%，Sharpe 1.23")
body("② MA Crossover：+1086%，Sharpe 0.98")
body("③ Momentum：+1056%，Sharpe 1.02")
body("")
body("📈 Triple MA 策略逻辑非常简单：")
body("短线 > 中线 > 长线 = 多头排列 = 买入")
body("短线跌破中线 = 卖出")
body("就这。")
body("")
body("🔗 我生成了分享链接，点进去可以看到完整结果：")
body("所有指标、收益曲线、交易明细都在里面")
body("")
body("👇 你也想测自己的策略？")
body("打开 quantflow.io 就行")
body("✅ 5 次免费回测/天")
body("✅ 12 种策略开箱即用")
body("✅ 输入 ticker 就能跑，30 秒出结果")
body("✅ 不用装任何东西，浏览器就行")
body("")
body("💬 评论区告诉我你的 ticker，我帮你跑")
divider()

# ═══════════════════════════════════════════════════════════
# 4. 标签
# ═══════════════════════════════════════════════════════════
h1("标签（复制粘贴到小红书标签栏）")

tag("#量化交易 #BTC #加密货币 #回测 #交易策略 #QuantFlow #免费工具 #TradingView")
divider()

# ═══════════════════════════════════════════════════════════
# 5. 图片顺序
# ═══════════════════════════════════════════════════════════
h1("图片发布顺序（6 张）")

pics = [
    ("图 1 · 封面", "BTC 收益曲线大图（Demo 页截图）—— 加文字标注 \"+1662%\" 吸引眼球"),
    ("图 2 · 指标特写", "指标卡片区域截图，用红圈标出 +1662% 和 Sharpe 1.23"),
    ("图 3 · 策略对比", "三行文字卡片：Triple MA vs MA Cross vs Momentum 的收益对比"),
    ("图 4 · 产品截图", "策略选择界面 —— 展示 12 种策略的分类排列"),
    ("图 5 · 结果展示", "结果页完整截图：指标卡 + 收益曲线 + 交易表"),
    ("图 6 · 行动召唤", "把 https://quantflow.pages.dev/demo 生成二维码，加文字\"扫码免费测试\""),
]

for title, desc in pics:
    h2(title)
    body(desc)

divider()

# ═══════════════════════════════════════════════════════════
# 6. 互动话术
# ═══════════════════════════════════════════════════════════
h1("评论区互动话术")

body("有人问「能帮我测一下 XXX 吗」→ 回复：")
body('')
body('    "好的！我帮你跑了一下 XXX 的 MA Cross 策略，')
body('     过去 3 年收益是 +XX%，这是完整结果链接 👇')
body('     [粘贴你的分享链接]"')
body('')
body('有人问「怎么用」→ 回复：')
body('')
body('    "直接打开 quantflow.io 就行，')
body('     输入 ticker → 选策略 → 点运行，')
body('     不用注册也能先看 Demo 👇')
body('     这是 Demo 页：[粘贴 Demo 链接]"')

divider()

# ═══════════════════════════════════════════════════════════
# 7. 发布检查清单
# ═══════════════════════════════════════════════════════════
h1("发布前检查清单")

checklist = [
    "✅ Demo 页面能正常打开",
    "✅ 分享链接有效且数据正确",
    "✅ 截图清晰，文字可读",
    "✅ 标签全部粘贴（至少 5 个）",
    "✅ 定位选「财经」「科技」相关",
    "✅ 发布时间：晚上 8-10 点或周末上午",
]
for item in checklist:
    bullet(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 8. 简易版（快速发布用）
# ═══════════════════════════════════════════════════════════
h1("简易版（如果上面太长，用这个）")

body("")
body("🚀 BTC 用 Triple MA 策略，5 年收益 +1662%")
body("")
body("📊 我做了个免费回测工具 QuantFlow，把 BTC 过去 5 年的数据跑了 12 种策略。")
body("Triple MA 拿了第一：+1662% 收益，Sharpe 1.23，33 笔交易，55% 胜率。")
body("")
body("🔗 完整结果看这里：[粘贴分享链接]")
body("")
body("👇 你也想跑自己的策略？")
body("去 quantflow.io，免费注册，每天 5 次回测，12 种策略开箱即用。")
body("输入 ticker 就能跑，不用写代码。")
body("")
body("💬 评论区告诉我你想测什么，我帮你跑")
tag("#量化交易 #BTC #回测 #交易策略 #QuantFlow")

# ═══════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════
output = "/root/quantflow/QuantFlow_小红书推广文案.docx"
doc.save(output)
print(f"✅ 已生成: {output}")
