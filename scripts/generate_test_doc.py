#!/usr/bin/env python3
"""Generate quantflow_manual_test_plan.docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ── Styles ───────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def checkbox(label):
    doc.add_paragraph(f"☐ {label}", style="List Bullet")

def info_box(title, content):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(content)

def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)
    return table

# ── Cover ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("QuantFlow 手动测试计划")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"版本 1.0 · {datetime.date.today().strftime('%Y-%m-%d')}").font.size = Pt(14)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("测试环境: https://quantflow-two.vercel.app\n")
info.add_run("后端 API: https://quantflow-v3q5.onrender.com\n")
info.add_run("预计耗时: 20 分钟")

doc.add_page_break()

# ── Test 1 ───────────────────────────────────────────────────────────────────
h1("测试 1: 落地页（2 分钟）")
h2("1.1 桌面端加载")
checkbox("打开 Chrome 无痕窗口 → https://quantflow-two.vercel.app")
checkbox("页面 3 秒内完成加载")
checkbox('大标题 "Backtest Any Trading Strategy in 5 Minutes" 可见')
checkbox('"Start for Free" 按钮可见且可点击')
checkbox("向下滚动，区块依次淡入（Framer Motion 动画）")
checkbox("策略卡片（MA Crossover / RSI / Bollinger）排列正常")
checkbox("定价卡片（Free / Pro $19 / Quant $49）3 列对齐")
checkbox("Footer 链接可见，版权年份正确")

h2("1.2 移动端布局")
checkbox("按 F12 → Toggle Device Toolbar → iPhone 12 Pro")
checkbox("标题字号自适应，不溢出")
checkbox("CTA 按钮堆叠为竖排")
checkbox("策略卡片单列排列")
checkbox("定价卡片单列排列")
checkbox("导航栏正常显示")

h2("1.3 导航跳转")
checkbox('点击 "Start for Free" → 跳转到 /register')

h2("Bug 记录")
make_table(["检查项", "通过", "问题描述"], [
    ["桌面端加载", "☐", ""],
    ["移动端布局", "☐", ""],
    ["导航跳转", "☐", ""],
])

doc.add_page_break()

# ── Test 2 ───────────────────────────────────────────────────────────────────
h1("测试 2: 注册与登录（3 分钟）")
h2("2.1 异常输入验证")

make_table(["#", "操作", "期望结果", "实际结果", "通过"], [
    ["1", "空表单直接提交", "表单验证拦截", "", "☐"],
    ["2", '输入 "abc" 到邮箱框', "Invalid email 格式提示", "", "☐"],
    ["3", '输入 "ab1" 到密码框', "密码至少 8 位", "", "☐"],
    ["4", '输入 "abcdefgh"', "密码必须包含数字", "", "☐"],
    ["5", '输入 "12345678"', "密码必须包含字母", "", "☐"],
    ["6", "输入已注册邮箱", "Email already exists", "", "☐"],
])

h2("2.2 正常注册")
para("测试账号: Email = mytest_xxx@gmail.com, Password = Test1234, Name = Manual Tester", bold=True)
checkbox("请求成功发出（Network 标签 POST /auth/register）")
checkbox("返回 201 + access_token + refresh_token")
checkbox("跳转到 /dashboard")
checkbox("侧边栏: Dashboard / Backtest / Billing / Settings")

h2("2.3 登录测试")
checkbox("正确密码 → 进入 Dashboard")
checkbox("错误密码 → 显示错误提示，不跳转，表单不清空")
checkbox("不存在的邮箱 → 显示错误提示")

h2("Bug 记录")
make_table(["检查项", "通过", "问题描述"], [
    ["异常输入验证", "☐", ""],
    ["正常注册", "☐", ""],
    ["登录测试", "☐", ""],
])

doc.add_page_break()

# ── Test 3 ───────────────────────────────────────────────────────────────────
h1("测试 3: Dashboard（1 分钟）")
checkbox("侧边栏 4 个导航项可见")
checkbox("统计卡片显示: Total Backtests / Win Rate / Avg Sharpe / Running Now")
checkbox('"No backtests yet" 空状态显示')
checkbox('点击 "Get started" → 跳转到 /backtest')
checkbox('点击 "New Backtest" 按钮 → 同样跳转')

make_table(["检查项", "通过", "问题描述"], [
    ["Dashboard 渲染", "☐", ""],
    ["导航跳转", "☐", ""],
])

doc.add_page_break()

# ── Test 4 ───────────────────────────────────────────────────────────────────
h1("测试 4: 回测配置向导（5 分钟）")
h2("4.1 Step 1: Ticker 模式")
checkbox("默认选中 'Enter Ticker' 标签")
checkbox('输入 AAPL → 自动大写')
checkbox("点击日期预设 '1Y' → 日期自动填充（start=1年前, end=今天）")
checkbox("日期选择器可手动修改")
checkbox("点击 'Next' → 进入 Step 2")
checkbox("步骤指示器: Step 1 显示绿色对勾 ✅")

h2("4.2 Step 2: 选择策略")
checkbox("3 张策略卡片显示，图标+标题+描述")
checkbox("点击 MA Crossover → 绿色边框 + 'Selected' 徽章")
checkbox("点击 RSI → 选中切换，上一个自动取消")
checkbox("点击 'Next' → 进入 Step 3")

h2("4.3 Step 3: 策略参数（动态渲染）")
para("MA Cross 参数:", bold=True)
checkbox("Fast Period 滑块: 拖动到 5 → 数字变为 5")
checkbox("Slow Period 滑块: 拖动到 50 → 数字变为 50")
checkbox("MA Type 切换: 点 EMA → 切换成功，SMA 取消高亮")

para("切回 Step 2 → 选 RSI:", bold=True)
checkbox("参数动态切换为 RSI Period / Oversold / Overbought")
checkbox("3 个滑块均正常拖动")

para("切回 Step 2 → 选 Bollinger:", bold=True)
checkbox("参数变为 Period + Std Dev Multiplier")
checkbox("Std Dev 滑块步长为 0.1")

h2("4.4 Step 4: 运行设置")
checkbox("Initial Capital 默认 $10,000")
checkbox("Commission 默认 0.10%")
checkbox("Backtest Name 可输入文字")
checkbox('"Free tier: N backtests remaining today" 提示可见')
checkbox("Configuration Summary 卡片摘要正确")

h2("4.5 步骤导航")
checkbox('点 "Back" 3 次 → 回到 Step 1')
checkbox('点 "Next" 3 次 → 到 Step 4')
checkbox("步骤指示器实时更新")

h2("4.6 CSV 上传模式")
para("上传 test_normal.csv (1000行):", bold=True)
checkbox("拖拽区域显示虚线边框+上传图标")
checkbox("上传后显示绿色 ✅ + 'Valid'")
checkbox("预览前 5 行表格可见")
checkbox("列名: date/open/high/low/close/volume")

para("上传 test_bad_data.csv (含异常):", bold=True)
checkbox("显示红色 ❌ + 'Issues found'")
checkbox("错误描述清晰")

para("上传 test_minimal.csv (61行):", bold=True)
checkbox("解析成功")
checkbox("仅 close 列也能正常识别")

make_table(["检查项", "通过", "问题描述"], [
    ["Ticker 输入", "☐", ""],
    ["策略选择", "☐", ""],
    ["参数动态渲染", "☐", ""],
    ["CSV 上传", "☐", ""],
    ["步骤导航", "☐", ""],
])

doc.add_page_break()

# ── Test 5 ───────────────────────────────────────────────────────────────────
h1("测试 5: 回测执行（2 分钟）")
h2("5.1 正常回测")
para("配置: Ticker=AAPL, Strategy=MA Cross (10/30), Capital=$10,000", bold=True)
checkbox('点击绿色 "Run Backtest" → 按钮变灰，显示 Loading 动画')
checkbox("Network 标签显示 POST /backtest/run-sync")
checkbox("完成后自动跳转到 /results/[id]")

h2("5.2 防重复提交")
checkbox("快速双击 Run 按钮 → 按钮立即变灰，不发第二个请求")

h2("5.3 预览面板")
checkbox("右侧面板显示策略标签（AAPL + MA Crossover）")
checkbox("演示曲线图可见（绿色SVG）")
checkbox("指标预览: Sharpe / Sortino / Max DD / Win Rate / Profit Factor / Total Trades")

make_table(["检查项", "通过", "问题描述"], [
    ["正常回测", "☐", ""],
    ["防重复提交", "☐", ""],
    ["预览面板", "☐", ""],
])

doc.add_page_break()

# ── Test 6 ───────────────────────────────────────────────────────────────────
h1("测试 6: 结果展示（3 分钟）")
h2("6.1 页面头部")
checkbox("面包屑导航: ← Backtest 链接")
checkbox("回测名称 + 日期范围 + 状态（Completed 绿色点）")
checkbox("操作按钮: Share / Export / New")

h2("6.2 指标卡片")
checkbox("6 张卡片 2×3 网格排列")
checkbox("Total Return: 正数绿色，负数红色")
checkbox("数字计数动画（从 0 增长到目标值）")
checkbox("vs B&H 对比显示（△ 箭头 + 百分比差异）")
checkbox("Max Drawdown 红色负数")
checkbox("Win Rate 正常 0~100%")

h2("6.3 图表交互")
h3("Equity Curve:")
checkbox("图表非空白，正常渲染")
checkbox("绿色线 = 策略收益，灰色虚线 = Buy & Hold")
checkbox("图例: 绿色块 Strategy / 灰色块 Buy & Hold")
checkbox("鼠标悬停 → 十字准线 + tooltip（日期/策略值/基准值）")
checkbox("滚轮缩放正常")
checkbox("鼠标拖拽平移正常")
checkbox("等比例缩放正常（窗口拉伸）")

h3("Drawdown 图表:")
checkbox("红色渐变填充区域")
checkbox("标题显示最大回撤值")

h2("6.4 策略信息面板")
checkbox("Symbol / Date Range / Strategy Name / Initial Capital")
checkbox("策略参数以独立标签展示")

h2("6.5 交易记录表格")
checkbox("列: # / Entry Date / Exit Date / Side / Entry $ / Exit $ / Return / P&L")
checkbox("盈利行绿色，亏损行红色")
checkbox("Side 列显示绿色 'LONG' 徽章")
checkbox("点击 Return 列表头 → 升序/降序切换（Chevron 图标）")
checkbox("点击 P&L 列表头 → 排序")
checkbox("分页: Prev / 页码 / Next 按钮")
checkbox("总数显示 'Showing X–Y of Z trades'")

h2("6.6 分享功能")
checkbox('点击 Share → 显示 "Copied" 对勾（2秒后消失）')
checkbox("生成的链接格式: /share/backtest/UUID")

make_table(["检查项", "通过", "问题描述"], [
    ["指标卡片", "☐", ""],
    ["图表渲染+交互", "☐", ""],
    ["交易表格", "☐", ""],
    ["分享功能", "☐", ""],
])

doc.add_page_break()

# ── Test 7 ───────────────────────────────────────────────────────────────────
h1("测试 7: Billing 页面（1 分钟）")
checkbox("侧边栏点 Billing → 页面正常加载")
checkbox("当前计划显示 'Free' + 'Inactive' 状态")
checkbox("月付/年付滑块切换 → 价格更新")
checkbox("年付模式显示 'Save 30%' 绿色徽章")
checkbox("Pro 卡片有 'Most Popular' 绿色徽章")
checkbox("Free 卡片按钮灰显 'Current Plan'")
checkbox("功能对比表 10 行 × 4 列显示正常")
checkbox("Pro 列绿色高亮")
checkbox("测试卡号区显示 4 张卡片信息")
checkbox("移动端: 表格可横向滚动")

make_table(["检查项", "通过", "问题描述"], [
    ["Billing 页面", "☐", ""],
])

doc.add_page_break()

# ── Test 8 ───────────────────────────────────────────────────────────────────
h1("测试 8: 错误处理（2 分钟）")
h2("8.1 网络断开")
checkbox("DevTools → Network → Offline")
checkbox('点击 "Run Backtest" → 显示错误提示（不白屏）')
checkbox("恢复 Online → 可重试")

h2("8.2 未登录拦截")
checkbox("无痕窗口直接访问 /dashboard → 重定向到登录页")
checkbox("直接访问 /backtest → 重定向或提示登录")

h2("8.3 刷新持久化")
checkbox("结果页按 F5 → 数据不丢失")
checkbox("回测配置页 F5 → 参数保留（localStorage）")

h2("8.4 冷启动处理")
checkbox("后端无请求 15+ 分钟后，首次请求显示 Loading")
checkbox("不白屏、不崩溃")

make_table(["检查项", "通过", "问题描述"], [
    ["网络断开", "☐", ""],
    ["未登录拦截", "☐", ""],
    ["刷新持久化", "☐", ""],
])

doc.add_page_break()

# ── Test 9 ───────────────────────────────────────────────────────────────────
h1("测试 9: 浏览器兼容性（3 分钟）")

make_table(["浏览器", "落地页", "注册登录", "回测配置", "回测执行", "结果展示", "整体评价"], [
    ["Chrome", "☐", "☐", "☐", "☐", "☐", ""],
    ["Firefox", "☐", "☐", "☐", "☐", "☐", ""],
    ["Safari (Mac)", "☐", "☐", "☐", "☐", "☐", ""],
    ["Edge", "☐", "☐", "☐", "☐", "☐", ""],
    ["iPhone Safari", "☐", "☐", "☐", "☐", "☐", ""],
    ["Android Chrome", "☐", "☐", "☐", "☐", "☐", ""],
])

make_table(["屏幕尺寸", "宽度", "落地页", "回测页", "结果页", "评价"], [
    ["桌面", "1920×1080", "☐", "☐", "☐", ""],
    ["笔记本", "1366×768", "☐", "☐", "☐", ""],
    ["iPad", "768×1024", "☐", "☐", "☐", ""],
    ["iPhone", "375×812", "☐", "☐", "☐", ""],
])

doc.add_page_break()

# ── Test 10 ──────────────────────────────────────────────────────────────────
h1("测试 10: 持久化与状态管理（1 分钟）")
checkbox("Step 1: 输入 AAPL, 选 1Y 预设")
checkbox("Step 2: 选 RSI 策略")
checkbox("Step 3: 调整 Oversold 为 25")
checkbox("Step 4: Initial Capital 改为 $50,000")
checkbox("F5 刷新页面 —— 所有设置保留")
checkbox("切换到 CSV 模式再切回来 —— Ticker 设置保留")
checkbox("登出后重新登录 —— 设置保留")

make_table(["检查项", "通过", "问题描述"], [
    ["F5 刷新后参数保留", "☐", ""],
    ["Tab 切换不丢数据", "☐", ""],
    ["登出登入保留", "☐", ""],
])

doc.add_page_break()

# ── Final Report ─────────────────────────────────────────────────────────────
h1("测试结果汇总")
make_table(["模块", "通过数", "失败数", "Bug 数", "状态"], [
    ["1. 落地页", "", "", "", "☐"],
    ["2. 注册登录", "", "", "", "☐"],
    ["3. Dashboard", "", "", "", "☐"],
    ["4. 回测配置", "", "", "", "☐"],
    ["5. 回测执行", "", "", "", "☐"],
    ["6. 结果展示", "", "", "", "☐"],
    ["7. Billing", "", "", "", "☐"],
    ["8. 错误处理", "", "", "", "☐"],
    ["9. 兼容性", "", "", "", "☐"],
    ["10. 持久化", "", "", "", "☐"],
    ["总计", "", "", "", ""],
])

doc.add_paragraph()
h2("最终判定")
checkbox("GREEN 🟢 — 全部通过，可推广获客")
checkbox("YELLOW 🟡 — 有非核心 Bug，修复后推广")
checkbox("RED 🔴 — 核心流程有阻塞性 Bug，必须先修复")

h2("发现的 Bug 清单")
doc.add_paragraph("（在此记录测试中发现的所有 Bug）")

h2("测试人签字")
doc.add_paragraph("测试人: __________    日期: __________    耗时: __________")

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "/root/quantflow/QuantFlow_Manual_Test_Plan.docx"
doc.save(output_path)
print(f"✅ Generated: {output_path}")
